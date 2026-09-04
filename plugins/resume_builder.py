"""
plugins/resume_builder.py
-------------------------
Tailored ATS-compliant resume and CV generator in Word (.docx) and Markdown.
Adapted from Awesome Claude Skills Tailored Resume Generator.
"""
from __future__ import annotations

import json
from pathlib import Path

PLUGIN = {
    "name": "resume_builder",
    "description": (
        "Generates ATS-optimized professional resumes and CVs tailored to specific job descriptions or roles. "
        "Outputs both high-fidelity Word document (.docx) and clean Markdown. "
        "Use when user asks to create, tailor, or update a resume or CV."
    ),
    "parameters": {
        "type": "OBJECT",
        "properties": {
            "target_role": {
                "type": "STRING",
                "description": "Target job title or role e.g. 'Senior AI Systems Engineer' or 'Data Scientist'",
            },
            "job_description": {
                "type": "STRING",
                "description": "Optional text of the job posting to tailor keywords for ATS optimization.",
            },
            "user_experience": {
                "type": "STRING",
                "description": "User background, past jobs, projects, education, and technical skills.",
            },
            "output_path": {
                "type": "STRING",
                "description": "Optional file path to save the resume (.docx or .md).",
            },
        },
        "required": ["target_role", "user_experience"],
    },
}


def run(parameters: dict, player=None, speak=None) -> str:
    role = parameters.get("target_role", "Software Engineer").strip()
    exp = parameters.get("user_experience", "").strip()
    jd = parameters.get("job_description", "").strip()
    out_path = parameters.get("output_path") or ""

    if speak:
        speak(f"Drafting tailored resume for {role}, sir.")

    try:
        from google import genai
        key_path = Path(__file__).resolve().parent.parent / "config" / "api_keys.json"
        api_key = ""
        if key_path.exists():
            api_key = json.loads(key_path.read_text(encoding="utf-8")).get("gemini_api_key", "")

        client = genai.Client(api_key=api_key)
        prompt = f"""You are an Executive Career Coach and ATS Resume Specialist.
Create an exceptional, impact-driven resume for the target role: '{role}'.

TARGET JOB DESCRIPTION:
{jd or 'Not provided - use industry-standard keywords'}

CANDIDATE EXPERIENCE & BACKGROUND:
{exp}

Rules:
1. Format with clean Markdown:
   # [Candidate Name]
   [Contact Info | LinkedIn | GitHub | City]

   ## Professional Summary
   Concise, high-impact 3-line value proposition.

   ## Core Technical Skills
   Grouped by category (Languages, Frameworks, Cloud/DevOps, AI/ML, Tools).

   ## Professional Experience
   Reverse chronological. Use XYZ formula: 'Accomplished [X] as measured by [Y], by doing [Z]'. Use strong action verbs.

   ## Key Projects
   Bullet points with architecture and measurable outcomes.

   ## Education & Certifications
2. Use clean bullet points and bold metrics.
"""
        response = client.models.generate_content(
            model="gemini-2.5-flash",
            contents=prompt,
        )
        resume_md = (response.text or "").strip()

        p = Path(out_path) if out_path else Path.home() / "Desktop" / f"Resume_{role[:20].replace(' ', '_')}.docx"
        if not p.is_absolute():
            p = Path.home() / "Desktop" / p.name
        p.parent.mkdir(parents=True, exist_ok=True)

        if p.suffix.lower() == ".docx":
            import docx
            doc = docx.Document()
            for line in resume_md.splitlines():
                ls = line.rstrip()
                if not ls: continue
                if ls.startswith("# "): doc.add_heading(ls[2:], level=1)
                elif ls.startswith("## "): doc.add_heading(ls[3:], level=2)
                elif ls.startswith("### "): doc.add_heading(ls[4:], level=3)
                elif ls.startswith("- ") or ls.startswith("* "): doc.add_paragraph(ls[2:], style='List Bullet')
                else: doc.add_paragraph(ls)
            doc.save(str(p))
        else:
            p.write_text(resume_md, encoding="utf-8")

        return f"📄 Tailored ATS resume successfully generated and saved to: {p.name} ({p.resolve()})"
    except Exception as e:
        return f"Resume generation failed: {e}"
