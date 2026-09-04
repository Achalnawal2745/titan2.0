"""
plugins/meeting_insights.py
---------------------------
Analyzes meeting transcripts, extracts decisions, action items with owners/deadlines,
and formats executive summaries.
Adapted from Awesome Claude Skills Meeting Insights Analyzer.
"""
from __future__ import annotations

import json
from pathlib import Path

PLUGIN = {
    "name": "meeting_insights",
    "description": (
        "Analyzes meeting audio/text transcripts or notes, extracts key decisions, creates "
        "action item tables with owners and deadlines, and drafts follow-up emails. "
        "Use when user asks to summarize a meeting, extract action items, or review call notes."
    ),
    "parameters": {
        "type": "OBJECT",
        "properties": {
            "transcript_text": {
                "type": "STRING",
                "description": "Raw meeting transcript or notes text.",
            },
            "transcript_file": {
                "type": "STRING",
                "description": "Optional path to a .txt or .docx file containing the transcript.",
            },
            "output_path": {
                "type": "STRING",
                "description": "Optional path to save the meeting summary (.md or .docx).",
            },
        },
        "required": [],
    },
}


def run(parameters: dict, player=None, speak=None) -> str:
    raw_text = parameters.get("transcript_text", "").strip()
    file_path = parameters.get("transcript_file", "").strip()
    out_path = parameters.get("output_path", "").strip()

    if not raw_text and file_path:
        p = Path(file_path)
        if p.exists():
            if p.suffix.lower() == ".docx":
                import docx
                d = docx.Document(str(p))
                raw_text = "\n".join(para.text for para in d.paragraphs if para.text.strip())
            else:
                raw_text = p.read_text(encoding="utf-8", errors="ignore")

    if not raw_text:
        return "Please provide meeting transcript text or a valid transcript file path."

    if speak:
        speak("Analyzing meeting transcript and extracting action items, sir.")

    try:
        from google import genai
        key_path = Path(__file__).resolve().parent.parent / "config" / "api_keys.json"
        api_key = ""
        if key_path.exists():
            api_key = json.loads(key_path.read_text(encoding="utf-8")).get("gemini_api_key", "")

        client = genai.Client(api_key=api_key)
        prompt = f"""You are an Executive Chief of Staff.
Analyze this meeting transcript and produce a high-fidelity summary:

TRANSCRIPT:
{raw_text[:30000]}

Format in clean Markdown:
# Meeting Summary & Key Decisions
- Executive 3-bullet overview
- Decisions Finalized

# Action Items & Owners
| Task / Deliverable | Owner | Deadline | Priority |
|---|---|---|---|
(Populate table with concrete owners and deliverables)

# Topics Discussed
(Detailed breakdown of agenda items and disagreements resolved)

# Draft Follow-up Email
Subject: [Clear Subject]
Body: [Professional email to attendees with recap and action items]
"""
        response = client.models.generate_content(
            model="gemini-2.5-flash",
            contents=prompt,
        )
        report_md = (response.text or "").strip()

        if out_path:
            out_p = Path(out_path)
            if not out_p.is_absolute():
                out_p = Path.home() / "Desktop" / out_p.name
            out_p.parent.mkdir(parents=True, exist_ok=True)
            if out_p.suffix.lower() == ".docx":
                import docx
                doc = docx.Document()
                for line in report_md.splitlines():
                    ls = line.rstrip()
                    if not ls: continue
                    if ls.startswith("# "): doc.add_heading(ls[2:], level=1)
                    elif ls.startswith("## "): doc.add_heading(ls[3:], level=2)
                    elif ls.startswith("### "): doc.add_heading(ls[4:], level=3)
                    elif ls.startswith("- ") or ls.startswith("* "): doc.add_paragraph(ls[2:], style='List Bullet')
                    else: doc.add_paragraph(ls)
                doc.save(str(out_p))
            else:
                out_p.write_text(report_md, encoding="utf-8")
            return f"📋 Meeting insights extracted and saved to: {out_p.name}\n\n{report_md[:1200]}..."

        return report_md
    except Exception as e:
        return f"Meeting analysis failed: {e}"
