"""
plugins/research_synthesizer.py
-------------------------------
Deep research, multi-source synthesis, and structured report generator with citations.
Adapted from Awesome Claude Skills Content Research Writer.
"""
from __future__ import annotations

import json
from pathlib import Path

PLUGIN = {
    "name": "research_synthesizer",
    "description": (
        "Conducts multi-source web research on complex topics and generates structured research reports "
        "with executive summaries, comparative tables, pros/cons, and citations. "
        "Use when user asks for deep research, competitor analysis, or market/technical surveys."
    ),
    "parameters": {
        "type": "OBJECT",
        "properties": {
            "topic": {
                "type": "STRING",
                "description": "The research question or topic to analyze.",
            },
            "output_path": {
                "type": "STRING",
                "description": "Optional file path to save the research report (.md or .docx).",
            },
            "focus_areas": {
                "type": "ARRAY",
                "description": "Specific areas to focus on (e.g. ['Performance', 'Pricing', 'Security'])",
                "items": {"type": "STRING"},
            },
        },
        "required": ["topic"],
    },
}


def run(parameters: dict, player=None, speak=None) -> str:
    topic = parameters.get("topic", "").strip()
    if not topic:
        return "Please provide a research topic."

    focus = parameters.get("focus_areas") or []
    out_path = parameters.get("output_path") or ""

    if speak:
        speak(f"Beginning research synthesis on {topic}, sir.")

    try:
        from google import genai
        # Load API key
        key_path = Path(__file__).resolve().parent.parent / "config" / "api_keys.json"
        api_key = ""
        if key_path.exists():
            api_key = json.loads(key_path.read_text(encoding="utf-8")).get("gemini_api_key", "")

        client = genai.Client(api_key=api_key)
        focus_str = f"Focus particularly on: {', '.join(focus)}" if focus else ""
        prompt = f"""You are a Principal Research Analyst.
Generate an exhaustive, highly structured research briefing on: {topic}
{focus_str}

Format in clean Markdown:
# Executive Summary
Key takeaways in 3-4 bullet points.

# Market & Technical Analysis
Detailed breakdown with data and specifics.

# Comparative Matrix
| Criteria | Approach / Solution A | Approach / Solution B | Winner / Recommendation |
|---|---|---|---|
(Provide rich tabular comparison)

# Strategic Implications & Next Steps
Actionable recommendations.

# References & Industry Standards
Formal citations and standard references.
"""
        response = client.models.generate_content(
            model="gemini-2.5-flash",
            contents=prompt,
        )
        report_md = (response.text or "").strip()

        if out_path:
            p = Path(out_path)
            if not p.is_absolute():
                p = Path.home() / "Desktop" / p.name
            p.parent.mkdir(parents=True, exist_ok=True)
            if p.suffix.lower() == ".docx":
                import docx
                doc = docx.Document()
                for line in report_md.splitlines():
                    ls = line.rstrip()
                    if not ls: continue
                    if ls.startswith("# "): doc.add_heading(ls[2:], level=1)
                    elif ls.startswith("## "): doc.add_heading(ls[3:], level=2)
                    elif ls.startswith("### "): doc.add_heading(ls[4:], level=3)
                    elif ls.startswith("- ") or ls.startswith("* "): doc.add_paragraph(ls[2:], style='List Bullet')
                    else: doc.add_paragraph(ls)
                doc.save(str(p))
            else:
                p.write_text(report_md, encoding="utf-8")
            return f"✅ Research briefing complete and saved to {p.name}:\n\n{report_md[:1500]}..."

        return report_md
    except Exception as e:
        return f"Research synthesis failed: {e}"
